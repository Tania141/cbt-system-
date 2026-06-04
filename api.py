"""
АКТ СИСТЕМ REST API v2.1 - Template-based document generation + Claude AI
"""
import os, io, re, json, base64
from datetime import datetime
from flask import Flask, request, jsonify
from flask_cors import CORS
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side
from docx import Document
import dropbox
from dropbox.exceptions import ApiError
from dropbox.files import WriteMode
import anthropic

app = Flask(__name__)
CORS(app, origins=[
    "https://tania141.github.io",
    "http://localhost",
    "http://127.0.0.1",
    "null"  # file:// local files
])

DROPBOX_TOKEN         = os.environ.get("DROPBOX_TOKEN", "")
DROPBOX_REFRESH_TOKEN = os.environ.get("DROPBOX_REFRESH_TOKEN", "")
DROPBOX_APP_KEY       = os.environ.get("DROPBOX_APP_KEY", "")
DROPBOX_APP_SECRET    = os.environ.get("DROPBOX_APP_SECRET", "")
DROPBOX_FOLDER        = os.environ.get("DROPBOX_FOLDER", "/AKT_Projects")
TEMPLATES_FOLDER      = os.environ.get("TEMPLATES_FOLDER", "/AKT_Templates")
ANTHROPIC_API_KEY     = os.environ.get("ANTHROPIC_API_KEY", "")

TEMPLATE_FILES = {
    "protokol2":   "Protokol_2_Template.docx",
    "protokol2p2": "Protokol_2_Part2_Template.docx",
    "protokol2a":  "Protokol_2a_Template.docx",
    "obrazec3":    "template_obrazec3.docx",
    "akt5":        "Akt_5_Template.docx",
    "akt6":        "Akt_6_Template.docx",
    "akt7":        "Akt_7_Template.docx",
    "akt8":        "Akt_8_Template.docx",
    "akt9":        "Akt_9_Beton_Dnevnik_Template.docx",
    "akt10":       "Akt_10_Template.docx",
    "akt11":       "Akt_11_Template.docx",
    "akt13":       "Akt_13_Template.docx",
    "akt14":       "template_akt14_1.docx",
    "zapovedna":   "Zapovedna_Template.docx",
    "akt15":       "Akt_15_Template.docx",
    "akt16":       "Akt_16_Template.docx",
    "doklad":      "Okonchatelen_Doklad_Template.docx",
}

DOC_LABELS = {
    "protokol2":   "Протокол_2",
    "protokol2p2": "Протокол_2_Част_2",
    "protokol2a":  "Протокол_2а",
    "obrazec3":    "Obrazec_3",
    "akt5":        "Akt_5",
    "akt6":        "Akt_6",
    "akt7":        "Akt_7",
    "akt8":        "Akt_8",
    "akt9":        "Akt_9_Beton_Dnevnik",
    "akt10":       "Akt_10",
    "akt11":       "Akt_11",
    "akt13":       "Akt_13",
    "akt14":       "Akt_14",
    "zapovedna":   "Zapovedna_Kniga",
    "akt15":       "Akt_15",
    "akt16":       "Akt_16",
    "doklad":      "Okonchatelen_Doklad",
}

# ── Dropbox helpers ───────────────────────────────────────────────────────────
def get_dropbox():
    if DROPBOX_REFRESH_TOKEN and DROPBOX_APP_KEY and DROPBOX_APP_SECRET:
        return dropbox.Dropbox(
            oauth2_refresh_token=DROPBOX_REFRESH_TOKEN,
            app_key=DROPBOX_APP_KEY,
            app_secret=DROPBOX_APP_SECRET
        )
    elif DROPBOX_TOKEN:
        return dropbox.Dropbox(DROPBOX_TOKEN)
    return None

def dbx_upload(dbx, path, data):
    dbx.files_upload(data, path, mode=WriteMode.overwrite, autorename=False, mute=True)

def dbx_download(dbx, path):
    _, res = dbx.files_download(path)
    return res.content

def dbx_create_folder(dbx, path):
    try:
        dbx.files_create_folder_v2(path)
    except ApiError as e:
        err_str = str(e)
        if "conflict" not in err_str and "already_exists" not in err_str:
            raise

def dbx_list_folder(dbx, path):
    try:
        r = dbx.files_list_folder(path)
        entries = list(r.entries)
        while r.has_more:
            r = dbx.files_list_folder_continue(r.cursor)
            entries.extend(r.entries)
        return entries
    except ApiError:
        return []

def get_shared_link(dbx, path):
    try:
        links = dbx.sharing_list_shared_links(path=path, direct_only=True).links
        if links:
            return links[0].url.replace("?dl=0", "?dl=1")
        return dbx.sharing_create_shared_link_with_settings(path).url.replace("?dl=0", "?dl=1")
    except Exception:
        return ""

# ── Excel helpers ─────────────────────────────────────────────────────────────
def build_passport_excel(rows):
    wb = Workbook()
    ws = wb.active
    ws.title = "Паспорт"
    ws.column_dimensions["A"].width = 38
    ws.column_dimensions["B"].width = 65
    hf  = Font(name="Calibri", bold=True, size=11, color="FFFFFF")
    hfill = PatternFill("solid", fgColor="1E3A5F")
    df  = Font(name="Calibri", size=11)
    af  = PatternFill("solid", fgColor="EAF1F8")
    thin = Side(style="thin", color="D0DBE6")
    brd  = Border(left=thin, right=thin, top=thin, bottom=thin)
    ws["A1"] = "Поле"; ws["B1"] = "Стойност"
    ws["A1"].font = hf; ws["B1"].font = hf
    ws["A1"].fill = hfill; ws["B1"].fill = hfill
    for i, row in enumerate(rows, start=2):
        if isinstance(row, (list, tuple)) and len(row) >= 2:
            ws.cell(i, 1, str(row[0])); ws.cell(i, 2, str(row[1]) if row[1] else "")
            ws.cell(i, 1).font = df; ws.cell(i, 2).font = df
            ws.cell(i, 1).border = brd; ws.cell(i, 2).border = brd
            if i % 2 == 0:
                ws.cell(i, 1).fill = af; ws.cell(i, 2).fill = af
    buf = io.BytesIO(); wb.save(buf); return buf.getvalue()

def read_passport_excel(data):
    ws = openpyxl.load_workbook(io.BytesIO(data)).active
    return {str(r[0]): str(r[1]) if r[1] else "" for r in ws.iter_rows(min_row=2, values_only=True) if r[0]}

def rows_to_dict(rows):
    return {str(r[0]): str(r[1]) if r[1] else "" for r in rows if isinstance(r, (list, tuple)) and len(r) >= 2}

# ── Placeholder logic ─────────────────────────────────────────────────────────
def two_names(s):
    parts = (s or "").strip().split()
    return " ".join(parts[:2])

def one_and_three(s):
    parts = (s or "").strip().split()
    if len(parts) >= 3:
        return f"{parts[0]} {parts[2]}"
    return " ".join(parts)

def fmt_date(v):
    if not v: return ""
    v = v.strip()
    if "." in v: return v
    try: return datetime.fromisoformat(v[:10]).strftime("%d.%m.%Y")
    except: return v

def extract_employees(d):
    n = int(d.get("Служители_Брой", 0) or 0)
    return [{
        "name": d.get(f"Служител_{i}_Име",""),
        "specialization": d.get(f"Служител_{i}_Специализация",""),
        "title": d.get(f"Служител_{i}_Титла","инж.")
    } for i in range(1, n+1) if d.get(f"Служител_{i}_Име","")]

def extract_projectants(d):
    n = int(d.get("Проектанти_Брой", 0) or 0)
    return [{
        "name": d.get(f"Проектант_{i}_Име",""),
        "ppp": d.get(f"Проектант_{i}_ППП",""),
        "specialization": d.get(f"Проектант_{i}_Специализация",""),
        "title": d.get(f"Проектант_{i}_Титла","инж.")
    } for i in range(1, n+1) if d.get(f"Проектант_{i}_Име","")]

def by_spec(items, spec):
    return [e for e in items if e.get("specialization") == spec]

def build_vazlogitel_block(d):
    tip   = d.get("Възложител_Тип", "Фирма")
    firma = d.get("Възложител_Фирма", "")
    adres = d.get("Възложител_Адрес", "")
    if tip in ("Физическо лице", "ФЛ"):
        return firma + (f", {adres}" if adres else "")
    eik  = d.get("Възложител_ЕИК", "")
    pred = d.get("Възложител_Представител", "")
    parts = [firma]
    if eik:   parts.append(f"ЕИК {eik}")
    if adres: parts.append(adres)
    if pred:  parts.append(f"представлявано от {pred}")
    return ", ".join(parts)

def build_projectants_list(projectants):
    lines = []
    for p in projectants:
        spec = p["specialization"]
        title = p["title"]
        name = p["name"]
        ppp = p["ppp"]
        kamara = "КАБ" if spec in ("Архитектура", "Паркоустройство и Благоустройство") else "КИИП"
        line = "Част " + spec + ": " + title + " " + name
        if ppp:
            line += ", рег. № " + ppp + " в " + kamara
        lines.append(line)
    return "\n".join(lines)

def build_employees_list(employees):
    return "\n".join(f"Част {e['specialization']}: {e['title']} {e['name']}" for e in employees)

def build_projectants_signatures(projectants):
    return "\n".join(
        f"Част {p['specialization']}: {p['title']} {p['name']} ….................................................................."
        for p in projectants
    )

def build_employees_signatures(employees):
    return "\n".join(
        f"Част {e['specialization']}: {e['title']} {e['name']} ….................................................................."
        for e in employees
    )

def build_placeholders(d):
    employees   = extract_employees(d)
    projectants = extract_projectants(d)

    geo     = (by_spec(employees, "Геодезия") or [{}])[0].get("name", d.get("Геодезист",""))
    sn_k    = (by_spec(employees, "Конструктивна") or [{}])[0].get("name", "")
    pj_k    = (by_spec(projectants, "Конструктивна") or [{}])[0].get("name", "")
    specs   = "; ".join(f"{e['title']} {e['name']} ({e['specialization']})" for e in employees) or d.get("Консултант_Управител","")
    vaz_tip = d.get("Възложител_Тип", "Фирма")
    upr     = d.get("Консултант_Управител", "")
    teh_ryk = d.get("Строител_ТехРък", "")
    str_upr = d.get("Строител_Управител", "")
    vaz_pr  = d.get("Възложител_Представител", "")

    return {
        "{{Строеж}}":                    d.get("Строеж",""),
        "{{Адрес}}":                     d.get("Адрес",""),
        "{{Консултант_Фирма}}":          d.get("Консултант_Фирма",""),
        "{{Консултант_ЕИК}}":            d.get("Консултант_ЕИК",""),
        "{{Консултант_Адрес}}":          d.get("Консултант_Адрес",""),
        "{{Консултант_Управител}}":      upr,
        "{{Консултант_Удостоверение}}":   d.get("Консултант_Удостоверение",""),
        "{{Управител_2имена}}":          two_names(upr),
        "{{Управител_1и3}}":             one_and_three(upr),
        "{{Строител_Фирма}}":            d.get("Строител_Фирма",""),
        "{{Строител_ЕИК}}":              d.get("Строител_ЕИК",""),
        "{{Строител_Адрес}}":            d.get("Строител_Адрес",""),
        "{{Строител_Управител}}":        str_upr,
        "{{Строител_Управител_2имена}}": two_names(str_upr),
        "{{Строител_Управител_1и3}}":    one_and_three(str_upr),
        "{{Строител_ТехРък}}":           teh_ryk,
        "{{ТехРък_2имена}}":             two_names(teh_ryk),
        "{{ТехРък_1и3}}":               one_and_three(teh_ryk),
        "{{tech_director}}":             teh_ryk,
        "{{Възложител_Тип}}":            vaz_tip,
        "{{Възложител_Фирма}}":          d.get("Възложител_Фирма",""),
        "{{Възложител_ЕИК}}":            d.get("Възложител_ЕИК","") if vaz_tip not in ("Физическо лице","ФЛ") else "",
        "{{Възложител_Адрес}}":          d.get("Възложител_Адрес",""),
        "{{Възложител_Представител}}":   vaz_pr if vaz_tip not in ("Физическо лице","ФЛ") else "",
        "{{Възложител_2имена}}":         two_names(vaz_pr),
        "{{Възложител_1и3}}":           one_and_three(vaz_pr),
        "{{Възложител_Блок}}":           build_vazlogitel_block(d),
        "{{РС_Номер}}":                  d.get("РС_Номер",""),
        "{{РС_Дата}}":                   fmt_date(d.get("РС_Дата","")),
        "{{РС_Издател}}":                d.get("РС_Издател",""),
        "{{РС_ВСила}}":                  fmt_date(d.get("РС_ВСила","")),
        "{{Геодезист}}":                 geo,
        "{{Геодезист_2имена}}":          two_names(geo),
        "{{Геодезист_1и3}}":            one_and_three(geo),
        "{{consultant_specialists}}":    specs,
        "{{constructor_name}}":          pj_k,
        "{{sn_konstruktivna}}":          sn_k,
        "{{СН_Конструктивна}}":          sn_k,
        "{{СН_Архитектура}}":            (by_spec(employees, "Архитектура") or [{}])[0].get("name", ""),
        "{{СН_Електро}}":                (by_spec(employees, "Електро") or [{}])[0].get("name", ""),
        "{{СН_ВиК}}":                    (by_spec(employees, "ВиК") or [{}])[0].get("name", ""),
        "{{СН_Геодезия}}":               (by_spec(employees, "Геодезия") or [{}])[0].get("name", ""),
        "{{СН_ПБ}}":                     (by_spec(employees, "ПБ") or [{}])[0].get("name", ""),
        "{{СН_Пътна}}":                  (by_spec(employees, "Пътна") or [{}])[0].get("name", ""),
        "{{СН_ОВК}}":                    (by_spec(employees, "ОВК и ЕЕ") or [{}])[0].get("name", ""),
        "{{pj_konstruktivna}}":          pj_k,
        "{{Проектанти_Списък}}":         build_projectants_list(projectants),
        "{{Консултанти_Списък}}":        build_employees_list(employees),
        "{{Проектанти_Подписи}}":        build_projectants_signatures(projectants),
        "{{Консултанти_Подписи}}":       build_employees_signatures(employees),
    }


# ── Template engine ───────────────────────────────────────────────────────────
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

def insert_paragraphs_after(para, lines, font_name="Times New Roman", font_size=12):
    from docx.shared import Pt
    ref = para._element
    parent = ref.getparent()
    idx = list(parent).index(ref)
    for i, line in enumerate(lines):
        new_p = OxmlElement("w:p")
        new_r = OxmlElement("w:r")
        new_rpr = OxmlElement("w:rPr")
        if para.runs:
            orig_rpr = para.runs[0]._r.find(qn("w:rPr"))
            if orig_rpr is not None:
                new_rpr = deepcopy(orig_rpr)
        new_r.append(new_rpr)
        new_t = OxmlElement("w:t")
        new_t.text = line
        new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        new_r.append(new_t)
        new_p.append(new_r)
        parent.insert(idx + 1 + i, new_p)

def replace_in_runs(para, replacements):
    full = "".join(r.text for r in para.runs)
    if not any(k in full for k in replacements):
        return False

    multiline_key = None
    multiline_val = None
    for k, v in replacements.items():
        if k in full and "\n" in v:
            multiline_key = k
            multiline_val = v
            break

    if multiline_key:
        lines = multiline_val.split("\n")
        first_line = full.replace(multiline_key, lines[0])
        for k, v in replacements.items():
            if k != multiline_key and "\n" not in v:
                first_line = first_line.replace(k, v)
        if para.runs:
            para.runs[0].text = first_line
            for r in para.runs[1:]:
                r.text = ""
        if len(lines) > 1:
            insert_paragraphs_after(para, lines[1:])
        return True

    new_text = full
    for k, v in replacements.items():
        new_text = new_text.replace(k, v)
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    return False

def fill_template(doc, replacements):
    paras = list(doc.paragraphs)
    for para in paras:
        replace_in_runs(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para, replacements)
    for section in doc.sections:
        for para in section.header.paragraphs:
            replace_in_runs(para, replacements)
        for para in section.footer.paragraphs:
            replace_in_runs(para, replacements)
    return doc

def generate_from_template(dbx, template_name, replacements):
    path = f"{TEMPLATES_FOLDER}/{template_name}"
    try:
        tpl_bytes = dbx_download(dbx, path)
    except ApiError:
        raise FileNotFoundError(f"Шаблонът '{template_name}' липсва в Dropbox → {path}")
    doc = Document(io.BytesIO(tpl_bytes))
    fill_template(doc, replacements)
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

# ── Routes ────────────────────────────────────────────────────────────────────
@app.route("/health", methods=["GET"])
def health():
    dbx = get_dropbox()
    dbx_ok = False
    if dbx:
        try: dbx.users_get_current_account(); dbx_ok = True
        except: pass

    templates = {}
    if dbx_ok:
        for key, fname in TEMPLATE_FILES.items():
            try:
                dbx.files_get_metadata(f"{TEMPLATES_FOLDER}/{fname}")
                templates[key] = f"OK ({fname})"
            except:
                templates[key] = f"ЛИПСВА ({fname})"

    return jsonify({
        "status": "ok", "version": "2.1",
        "timestamp": datetime.utcnow().isoformat(),
        "dropbox": "свързан" if dbx_ok else "не е свързан",
        "dropbox_folder": DROPBOX_FOLDER,
        "templates_folder": TEMPLATES_FOLDER,
        "templates": templates,
        "ai": "конфигуриран" if ANTHROPIC_API_KEY else "не е конфигуриран",
    })

@app.route("/api/passports", methods=["GET"])
def list_passports():
    dbx = get_dropbox()
    if not dbx: return jsonify({"error": "Dropbox не е конфигуриран"}), 503
    try:
        entries = dbx_list_folder(dbx, DROPBOX_FOLDER)
        result = []
        for e in entries:
            if hasattr(e, "path_display") and not e.path_display.endswith(".xlsx"):
                try:
                    data = dbx_download(dbx, f"{e.path_display}/passport.xlsx")
                    p = read_passport_excel(data)
                    result.append({"pi": p.get("ПИ",""), "stroej": p.get("Строеж",""),
                                   "address": p.get("Адрес",""), "consultant": p.get("Консултант_Фирма","")})
                except: pass
        return jsonify({"passports": result, "count": len(result)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/passports/<pi>", methods=["GET"])
def get_passport(pi):
    dbx = get_dropbox()
    if not dbx: return jsonify({"error": "Dropbox не е конфигуриран"}), 503
    try:
        data = dbx_download(dbx, f"{DROPBOX_FOLDER}/PI-{pi}/passport.xlsx")
        return jsonify({"pi": pi, "passport": read_passport_excel(data)})
    except ApiError:
        return jsonify({"error": f"Паспорт PI-{pi} не е намерен"}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/passports/<pi>", methods=["POST"])
def save_passport(pi):
    dbx = get_dropbox()
    if not dbx: return jsonify({"error": "Dropbox не е конфигуриран"}), 503
    body = request.get_json()
    if not body or "passport" not in body:
        return jsonify({"error": "Липсва поле 'passport'"}), 400
    folder = f"{DROPBOX_FOLDER}/PI-{pi}"
    path   = f"{folder}/passport.xlsx"
    try:
        dbx_create_folder(dbx, folder)
        dbx_upload(dbx, path, build_passport_excel(body["passport"]))
        return jsonify({"status": "ok", "pi": pi, "path": path})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/generate/<doc_type>", methods=["POST"])
def generate_document(doc_type):
    if doc_type not in TEMPLATE_FILES:
        return jsonify({"error": f"Непознат тип: {doc_type}. Позволени: {list(TEMPLATE_FILES.keys())}"}), 400
    dbx = get_dropbox()
    if not dbx: return jsonify({"error": "Dropbox не е конфигуриран"}), 503
    body = request.get_json()
    if not body or "passport" not in body:
        return jsonify({"error": "Липсва поле 'passport'"}), 400

    d    = rows_to_dict(body["passport"])
    pi   = str(body.get("pi", "unknown"))
    repl = build_placeholders(d)

    if doc_type == "zapovedna":
        repl["{{Заповедна_Номер}}"] = body.get("zapovedna_number", "___")
        repl["{{Заповедна_Дата}}"]  = fmt_date(body.get("zapovedna_date", ""))

    filename = f"PI-{pi}_{DOC_LABELS[doc_type]}.docx"
    folder   = f"{DROPBOX_FOLDER}/PI-{pi}"
    path     = f"{folder}/{filename}"

    try:
        doc_bytes = generate_from_template(dbx, TEMPLATE_FILES[doc_type], repl)
        dbx_create_folder(dbx, folder)
        dbx_upload(dbx, path, doc_bytes)
        file_url = get_shared_link(dbx, path)
        return jsonify({"status": "ok", "doc_type": doc_type, "pi": pi,
                        "filename": filename, "path": path, "file_url": file_url})
    except FileNotFoundError as e:
        return jsonify({"error": str(e),
                        "hint": f"Качи шаблона в Dropbox → {TEMPLATES_FOLDER}/"}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── Cloud Sync ────────────────────────────────────────────────────────────────
@app.route("/api/cloud/save", methods=["POST"])
def cloud_save():
    dbx = get_dropbox()
    if not dbx:
        return jsonify({"error": "Dropbox не е свързан"}), 503
    body = request.get_json()
    if not body:
        return jsonify({"error": "Липсва тяло"}), 400
    path = f"{DROPBOX_FOLDER}/_akt_projects.json"
    try:
        data = json.dumps(body, ensure_ascii=False, indent=2).encode('utf-8')
        dbx_upload(dbx, path, data)
        return jsonify({"status": "ok", "path": path})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/cloud/load", methods=["GET"])
def cloud_load():
    dbx = get_dropbox()
    if not dbx:
        return jsonify({"error": "Dropbox не е свързан"}), 503
    path = f"{DROPBOX_FOLDER}/_akt_projects.json"
    try:
        data = dbx_download(dbx, path)
        payload = json.loads(data.decode('utf-8'))
        return jsonify(payload)
    except Exception:
        return jsonify({"projects": [], "consultants": [], "history": []}), 200


# ── AI Протоколи ──────────────────────────────────────────────────────────────
@app.route("/api/ai/generate", methods=["POST"])
def ai_generate():
    """
    Приема prompt + (опционално) PDF файлове като base64,
    извиква Claude API и връща генерирания текст.

    Тяло (JSON):
    {
        "prompt": "...",           # задължително — пълният prompt
        "files": [                 # опционално — прикачени документи
            {
                "name": "RS.pdf",
                "data": "<base64>",
                "media_type": "application/pdf"
            }
        ]
    }
    """
    if not ANTHROPIC_API_KEY:
        return jsonify({"error": "ANTHROPIC_API_KEY не е конфигуриран в Railway"}), 503

    body = request.get_json()
    if not body or "prompt" not in body:
        return jsonify({"error": "Липсва поле 'prompt'"}), 400

    prompt = body["prompt"]
    files  = body.get("files", [])

    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

        # Изграждаме съдържанието на съобщението
        content = []

        # Добавяме прикачените PDF файлове
        for f in files:
            if f.get("media_type") == "application/pdf" and f.get("data"):
                content.append({
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": f["data"]
                    }
                })

        # Добавяме prompt-а като текст
        content.append({
            "type": "text",
            "text": prompt
        })

        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[{"role": "user", "content": content}]
        )

        result_text = "".join(
            block.text for block in response.content
            if hasattr(block, "text")
        )

        return jsonify({
            "status": "ok",
            "result": result_text,
            "input_tokens":  response.usage.input_tokens,
            "output_tokens": response.usage.output_tokens,
        })

    except anthropic.APIError as e:
        return jsonify({"error": f"Claude API грешка: {str(e)}"}), 502
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    print(f"АКТ СИСТЕМ API v2.1 | порт {port} | шаблони: {TEMPLATES_FOLDER}")
    app.run(host="0.0.0.0", port=port)
